import os
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
import natsort
from PIL import Image

# --- 用户配置 ---

# 1. 文件夹和文件名设置
IMAGE_FOLDER = '.'
OUTPUT_DOCX_NAME = '生成的图片文档_Final_NoBold.docx' # 更新了输出文件名
ALLOWED_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')

# 2. Word文档格式设置
IMAGE_WIDTH_IN_DOC = Inches(6.0)

# 3. 图注文字内容
FIGURE_PREFIX = "Fig. S19"
DESCRIPTION_BODY = ". Robustness of ancestral hybridization signals to outgroup choice, using Clematis loureiriana. This figure displays the filtered ancestral hybridization signals for internal nodes, recalculated using C. loureiriana as the outgroup. The key hybridization signals for Clematis sect. Fruticella and sect. Meclatis are consistently recovered."

# 4. 图片压缩配置
RESIZE_IMAGES = True
MAX_WIDTH_PIXELS = 1500
JPEG_QUALITY = 85

# --- 脚本主程序 ---

def add_images_to_word(folder_path, docx_name):
    """
    扫描文件夹中的图片，按需压缩后，添加带格式的描述并插入到Word文档中。
    """
    print(f"开始处理文件夹: {os.path.abspath(folder_path)}")

    try:
        files = os.listdir(folder_path)
        sorted_files = natsort.natsorted(files)
    except FileNotFoundError:
        print(f"错误：找不到文件夹 '{folder_path}'。")
        return

    doc = Document()
    image_count = 0
    page_counter = 1

    for filename in sorted_files:
        if filename.lower().endswith(ALLOWED_EXTENSIONS):
            try:
                image_path = os.path.join(folder_path, filename)
                print(f"正在处理图片: {filename} (Page {page_counter})")

                image_source = image_path

                if RESIZE_IMAGES:
                    with Image.open(image_path) as img:
                        if img.width > MAX_WIDTH_PIXELS:
                            print(f"  > 正在压缩图片... (原始尺寸: {img.width}x{img.height})")
                            aspect_ratio = img.height / img.width
                            new_height = int(MAX_WIDTH_PIXELS * aspect_ratio)
                            resized_img = img.resize((MAX_WIDTH_PIXELS, new_height), Image.Resampling.LANCZOS)
                            
                            img_stream = io.BytesIO()
                            
                            if img.format == 'JPEG':
                                resized_img.save(img_stream, format='JPEG', quality=JPEG_QUALITY)
                            elif img.format == 'PNG':
                                resized_img.save(img_stream, format='PNG', optimize=True)
                            else:
                                resized_img.save(img_stream, format=img.format)
                            
                            img_stream.seek(0)
                            image_source = img_stream
                            print(f"  > 压缩完成 (新尺寸: {MAX_WIDTH_PIXELS}x{new_height})")

                full_text = f"{FIGURE_PREFIX} (page {page_counter}){DESCRIPTION_BODY}(sample of {filename})"
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(full_text)
                
                # run.bold = True  # <--- 此行已移除，取消加粗

                font = run.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                font.color.rgb = RGBColor(0, 0, 0)
                
                doc.add_paragraph()
                doc.add_picture(image_source, width=IMAGE_WIDTH_IN_DOC)
                doc.add_page_break()

                image_count += 1
                page_counter += 1

            except Exception as e:
                print(f"处理图片 '{filename}' 时发生严重错误: {e}")

    if image_count > 0:
        try:
            doc.save(docx_name)
            print(f"\n处理完成！共 {image_count} 张图片已成功插入到 '{docx_name}' 文件中。")
        except Exception as e:
            print(f"保存Word文档时发生错误: {e}")
    else:
        print("\n在指定文件夹中没有找到任何符合条件的图片文件。")

if __name__ == '__main__':
    add_images_to_word(IMAGE_FOLDER, OUTPUT_DOCX_NAME)